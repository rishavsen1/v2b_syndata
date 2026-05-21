"""Render users_cars_sessions.docx from the short-overview content.

Two-page Word doc summarising the users → cars → sessions pipeline.
Style: 11pt Calibri body, blue (#1f4e79) headers, monospace code,
light grey shaded code/math snippets, orange callout sidebar for the
worked example.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

OUT = Path(__file__).parent / "users_cars_sessions.docx"
FIG_DAG = Path(__file__).parent / "figures" / "01_three_csvs_dag.png"

BLUE = RGBColor(0x1F, 0x4E, 0x79)
ORANGE = RGBColor(0xD8, 0x85, 0x3B)
LIGHT_BLUE_FILL = "DEEAF6"  # light blue header row fill
LIGHT_GREY_FILL = "F5F5F5"  # code background


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def _shade(cell, fill_hex: str) -> None:
    """Set a table cell background colour."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _para_shade(paragraph, fill_hex: str) -> None:
    """Set a paragraph background (light grey for code-like blocks)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _set_para_border(paragraph, side: str, color_hex: str, size_pt: int = 12) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bdr = OxmlElement(f"w:{side}")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), str(size_pt * 2))
    bdr.set(qn("w:space"), "4")
    bdr.set(qn("w:color"), color_hex)
    pBdr.append(bdr)


def set_section_margins(section, *, top=0.8, bottom=0.8, left=0.85, right=0.85) -> None:
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def title_paragraph(doc: Document, text: str, size: int = 18) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = BLUE
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(2)


def subtitle_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(6)


def section_heading(doc: Document, text: str, size: int = 12) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = BLUE
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def body_paragraph(doc: Document, *runs, size: int = 10):
    """runs: list of (text, {"bold":bool,"italic":bool,"mono":bool,"color":RGBColor})."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    for text, opts in runs:
        run = p.add_run(text)
        run.font.size = Pt(opts.get("size", size))
        run.font.bold = opts.get("bold", False)
        run.font.italic = opts.get("italic", False)
        if opts.get("mono"):
            run.font.name = "Consolas"
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                rFonts.set(qn(attr), "Consolas")
        else:
            run.font.name = "Calibri"
        if "color" in opts:
            run.font.color.rgb = opts["color"]
    return p


def numbered_step(doc: Document, n: int, head: str, body_runs: list) -> None:
    """Step paragraph: '1. **head** — body...' with monospace where needed."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)

    run = p.add_run(f"{n}. ")
    run.font.size = Pt(9.5)
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.color.rgb = BLUE

    run = p.add_run(head)
    run.font.size = Pt(9.5)
    run.font.bold = True
    run.font.name = "Calibri"

    if body_runs:
        run = p.add_run(" — ")
        run.font.size = Pt(9.5)
        run.font.name = "Calibri"

    for txt, opts in body_runs:
        run = p.add_run(txt)
        run.font.size = Pt(opts.get("size", 9.5))
        if opts.get("mono"):
            run.font.name = "Consolas"
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement("w:rFonts")
            for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                rFonts.set(qn(attr), "Consolas")
            rPr.append(rFonts)
        else:
            run.font.name = "Calibri"
        if opts.get("italic"):
            run.font.italic = True
        if opts.get("bold"):
            run.font.bold = True


def styled_table(doc: Document, header: list[str], rows: list[list[str]],
                 widths_in: list[float] | None = None, font_size: float = 9.0) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.autofit = False
    if widths_in:
        for i, w in enumerate(widths_in):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    # header row
    for j, h in enumerate(header):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = BLUE
        run.font.name = "Calibri"
        _shade(cell, LIGHT_BLUE_FILL)
    # body rows
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(font_size)
            run.font.name = "Calibri"
    # tighter row heights
    for row in table.rows:
        tr = row._tr
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = OxmlElement("w:trPr")
            tr.insert(0, trPr)
        h = OxmlElement("w:trHeight")
        h.set(qn("w:val"), "0")
        h.set(qn("w:hRule"), "atLeast")
        trPr.append(h)
    # apply table style for borders
    table.style = "Table Grid"


def callout_box(doc: Document, lines: list[tuple[str, dict]]) -> None:
    """Worked-example sidebar — orange left border, light grey fill."""
    for i, (text, opts) in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.12)
        _para_shade(p, "FFF3E0")
        _set_para_border(p, "left", "D8853B", size_pt=12)
        if i == 0:
            _set_para_border(p, "top", "FFE0B2", size_pt=4)
        if i == len(lines) - 1:
            _set_para_border(p, "bottom", "FFE0B2", size_pt=4)

        run = p.add_run(text)
        run.font.size = Pt(opts.get("size", 9.0))
        if opts.get("bold"):
            run.font.bold = True
        if opts.get("italic"):
            run.font.italic = True
        if opts.get("mono"):
            run.font.name = "Consolas"
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement("w:rFonts")
            for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                rFonts.set(qn(attr), "Consolas")
            rPr.append(rFonts)
        else:
            run.font.name = "Calibri"
        if opts.get("color"):
            run.font.color.rgb = opts["color"]


def mono(txt: str, **extra) -> tuple[str, dict]:
    d = {"mono": True}
    d.update(extra)
    return (txt, d)


def plain(txt: str, **extra) -> tuple[str, dict]:
    return (txt, extra)


# ---------------------------------------------------------------------------
# build document
# ---------------------------------------------------------------------------

def build() -> None:
    doc = Document()

    # default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    for section in doc.sections:
        set_section_margins(section)

    # -- header ----------------------------------------------------------
    title_paragraph(doc, "V2B Synthetic Data — Users, Cars, Sessions Generation")
    subtitle_paragraph(doc, "How three CSVs encode user behavior, fleet physics, and per-day arrivals.")

    # -- Section 1: three CSVs (with figure) ---------------------------
    section_heading(doc, "1. The three CSVs")
    body_paragraph(doc,
                   plain("users.csv", mono=True),
                   plain(", "),
                   plain("cars.csv", mono=True),
                   plain(", and "),
                   plain("sessions.csv", mono=True),
                   plain(" form a causal chain — generated "),
                   plain("in this order", italic=True),
                   plain(": users first, then cars, then sessions. Each "),
                   plain("users.csv", mono=True),
                   plain(" row maps 1:1 to a "),
                   plain("cars.csv", mono=True),
                   plain(" row (same "),
                   plain("car_id", mono=True),
                   plain("); each "),
                   plain("car_id", mono=True),
                   plain(" produces 0..N sessions across the sim window."))

    if FIG_DAG.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(FIG_DAG), width=Inches(3.8))

    # -- Section 2: behavioural axes -----------------------------------
    section_heading(doc, "2. Three behavioral axes (per car)")
    styled_table(
        doc,
        ["Axis", "Meaning", "Range", "Where it bites"],
        [
            ["φ (phi) — frequency",   "P(show up on any given day)",
             "[0.0, 1.0]",            "Bernoulli gate, sessions step 2"],
            ["κ (kappa) — consistency","How regular arrival timing is",
             "[0.0, 1.0]",            "σ_eff = σ · (1 − κ · 0.5),  step 3"],
            ["δ (delta_km) — commute","One-way commute proxy",
             "[3, 100] km",           "shift_eff = shift − 0.003·δ, step 5"],
        ],
        widths_in=[1.5, 2.1, 0.9, 2.5],
        font_size=9.0,
    )

    # -- Section 3: region grid -----------------------------------------
    section_heading(doc, "3. Default region grid (5 regions)")
    styled_table(
        doc,
        ["Region", "φ", "κ", "δ (km)", "Wt", "Mental model"],
        [
            ["stable_commuter",    "[0.85, 1.00]", "[0.75, 1.00]", "[40, 80]",  "0.35", "Long-distance office, daily"],
            ["flexible_local",     "[0.70, 0.95]", "[0.50, 0.80]", "[5, 15]",   "0.25", "Local, frequent + flexible"],
            ["irregular_distant",  "[0.40, 0.70]", "[0.20, 0.50]", "[40, 100]", "0.20", "Long commute, ~3 days/wk"],
            ["occasional_visitor", "[0.05, 0.20]", "[0.10, 0.40]", "[3, 50]",   "0.10", "Rare drop-in"],
            ["erratic",            "[0.30, 0.70]", "[0.05, 0.30]", "[5, 80]",   "0.10", "Unpredictable schedule"],
        ],
        widths_in=[1.45, 1.05, 1.05, 0.95, 0.45, 2.05],
        font_size=8.5,
    )

    # ===================================================================
    # PAGE 2
    # ===================================================================
    doc.add_page_break()

    # -- Section 4: users.csv ------------------------------------------
    section_heading(doc, "4. users.csv generation")
    numbered_step(doc, 1, "Optional Dirichlet perturb",
                  [plain("if "),
                   mono("axes_distribution_dirichlet_alpha < 1e6"),
                   plain(", "),
                   mono("realized_weights ~ Dirichlet(weights · α)"),
                   plain("; else verbatim. Logged to manifest.")])
    numbered_step(doc, 2, "Assign region", [mono("region ~ Categorical(realized_weights)")])
    numbered_step(doc, 3, "Sample axes (independent in 3D box)",
                  [mono("φ ~ U(freq_lo, hi)"), plain(", "),
                   mono("κ ~ U(consist_lo, hi)"), plain(", "),
                   mono("δ ~ U(dist_km_lo, hi)")])
    numbered_step(doc, 4, "Assign negotiation type",
                  [mono("negotiation_type ~ Categorical(negotiation_mix)"),
                   plain(" — 4 CONSENT clusters (I/II/III/IV)")])
    numbered_step(doc, 5, "Sample CONSENT weights",
                  [mono("(w1, w2) ~ N(cluster_μ, cluster_σ)"),
                   plain(", clipped ≥ 0, × "), mono("w_multiplier")])
    body_paragraph(doc,
                   plain("Schema: ", bold=True),
                   mono("car_id, region, phi, kappa, delta_km, negotiation_type, w1, w2"))

    # -- Section 5: cars.csv -------------------------------------------
    section_heading(doc, "5. cars.csv generation")
    numbered_step(doc, 1, "Branch by battery_heterogeneity",
                  [plain("homog → all "), mono("argmax(battery_mix)"),
                   plain("; mixed α ≥ 1e6 → per car "), mono("Categorical(battery_mix)"),
                   plain("; mixed α < 1e6 → once-per-sample "),
                   mono("realized_mix ~ Dirichlet(battery_mix · α)"),
                   plain(", then per car Categorical.")])
    numbered_step(doc, 2, "Lookup BATTERY_SPECS",
                  [mono("capacity_kwh ∈ {24, 40, 75, 100}"),
                   plain("; SoC bounds per class (typ. [10, 100]).")])
    body_paragraph(doc,
                   plain("Schema: ", bold=True),
                   mono("car_id, capacity_kwh, min_allowed_soc, max_allowed_soc, battery_class"))

    # -- Section 6: sessions.csv (10 steps) ----------------------------
    section_heading(doc, "6. sessions.csv generation (10 steps)")
    body_paragraph(doc,
                   plain("For each ", italic=True),
                   plain("car_id", mono=True, italic=True),
                   plain(" × each weekday in sim_window — rejection sampling, max 8 retries.",
                         italic=True))

    numbered_step(doc, 1, "Resolve region", [plain("from "), mono("users.csv")])
    numbered_step(doc, 2, "Attendance gate",
                  [mono("draw ~ U(0,1)"), plain("; skip day if "),
                   mono("draw > φ")])
    numbered_step(doc, 3, "Bivariate Gaussian copula sample",
                  [plain("on "),
                   mono("(arrival_hour, dwell_hours)"),
                   plain("; marginals "),
                   mono("TruncNorm(μ_r, σ_eff; [lo, hi])"),
                   plain(" × "),
                   mono("Weibull(k_r, λ_r)"),
                   plain(". "),
                   mono("(u_arr, u_dwell) ~ N₂(0, [[1, ρ], [ρ, 1]])"),
                   plain("; "),
                   mono("σ_eff = σ_r · (1 − κ · 0.5)"),
                   plain("; "),
                   mono("ρ = region.copula.rho_gaussian"),
                   plain(" — typ. negative ⇒ early arrivers stay longer. "
                         "Departure = arrival + dwell.")])
    numbered_step(doc, 4, "Snap arrival", [plain("to 15-min grid")])
    numbered_step(doc, 5, "Arrival SoC",
                  [mono("shift_eff = shift_r − 0.003 · δ"), plain("; "),
                   mono("arrival_soc ~ Beta(α_r, β_r) + shift_eff")])
    numbered_step(doc, 6, "Required SoC at departure",
                  [mono("required_soc ~ TruncNorm(85, 5)"),
                   plain(", clamped to "),
                   mono("[max(min_depart_soc, arrival_soc+ε), max_allowed_soc]")])
    numbered_step(doc, 7, "D5 reachability",
                  [mono("energy_needed = (required − arrival)/100 · capacity"),
                   plain("; "),
                   mono("max_charge = rate · dwell · 0.96"),
                   plain("; retry from step 3 if "),
                   mono("energy_needed > max_charge"),
                   plain(" (≤ 8x).")])
    numbered_step(doc, 8, "Floor duration", [plain("to 15-min multiples")])
    numbered_step(doc, 9, "Non-overlap (C7)",
                  [plain("drop if overlaps prior session for same "),
                   mono("car_id"), plain(" today")])
    numbered_step(doc, 10, "Track external SoC",
                  [mono("previous_day_external_use_soc"),
                   plain(" = SoC delta since last building session")])
    body_paragraph(doc,
                   plain("Schema: ", bold=True),
                   mono("session_id, car_id, building_id, arrival, departure, "),
                   mono("duration_sec, arrival_soc, required_soc_at_depart, "),
                   mono("previous_day_external_use_soc"))

    # -- Section 7: worked example (callout) ---------------------------
    section_heading(doc, "7. Worked example")
    callout_box(doc, [
        ("car_id = 42, day = 2024-04-08 (Mon)",
         {"bold": True, "size": 9.5, "color": ORANGE}),
        ("users.csv:  region=flexible_local  φ=0.84  κ=0.62  δ=12.4 km  negotiation_type=II  w1=0.71  w2=0.48",
         {"mono": True, "size": 8.5}),
        ("cars.csv:   battery_class=bolt_40  capacity_kwh=40  min/max_soc=10/100",
         {"mono": True, "size": 8.5}),
        ("region:     f_arr=N(9.5, σ=1.5)  f_dwell=Weibull(1.8, 6.5)  f_soc=Beta(5,5)  shift=0  ρ=−0.2",
         {"mono": True, "size": 8.5}),
        ("Attendance: draw=0.31 < 0.84 → continue.",
         {"mono": True, "size": 8.5}),
        ("Copula:     σ_eff = 1.5·(1−0.62·0.5) = 1.035;  (u_arr, u_dwell) = (−0.42, 0.18)",
         {"mono": True, "size": 8.5}),
        ("            arrival_hour = 9.18  dwell = 7.6 h    → snap arrival = 09:15",
         {"mono": True, "size": 8.5}),
        ("Arrival SoC: shift_eff = −0.003·12.4 = −0.037;  Beta(5,5)=0.512 + shift = 0.475 → 47.5%",
         {"mono": True, "size": 8.5}),
        ("Required:    floor=max(0.80, 0.475+ε)=0.80;  TruncNorm(85,5) → 84.3%",
         {"mono": True, "size": 8.5}),
        ("D5:          need = (84.3−47.5)/100·40 = 14.7 kWh;  max = 20·7.6·0.96 = 146 kWh   ✓",
         {"mono": True, "size": 8.5}),
        ("Duration:    7.6 h → 7.5 h = 27000 sec",
         {"mono": True, "size": 8.5}),
        ("→ sessions.csv: session_id=412, arrival=2024-04-08 09:15, departure=16:45, duration_sec=27000,",
         {"mono": True, "size": 8.5, "bold": True}),
        ("              arrival_soc=47.5, required_soc_at_depart=84.3, prev_day_ext_use_soc=0.0",
         {"mono": True, "size": 8.5, "bold": True}),
    ])

    # -- Section 8: knob cheat sheet -----------------------------------
    section_heading(doc, "8. Knob cheat sheet")
    styled_table(
        doc,
        ["Bucket", "Knob", "Effect"],
        [
            ["ev_fleet",       "ev_count",                              "# car_ids in users.csv and cars.csv"],
            ["ev_fleet",       "battery_mix",                           "Simplex over leaf_24/bolt_40/m3_75/rivian_100"],
            ["ev_fleet",       "battery_heterogeneity",                 "homog vs mixed branch in cars.py"],
            ["ev_fleet",       "battery_mix_dirichlet_alpha",           "Per-sample mix perturbation (1e6 = off)"],
            ["user_behavior",  "axes_distribution",                     "5-region grid: defines (φ,κ,δ) bounds + weight"],
            ["user_behavior",  "negotiation_mix",                       "CONSENT cluster distribution"],
            ["user_behavior",  "w_multiplier",                          "Scales (w1, w2)"],
            ["user_behavior",  "min_depart_soc",                        "Floor on required_soc_at_depart"],
            ["user_behavior",  "axes_distribution_dirichlet_alpha",     "Per-sample region-weight perturbation"],
            ["user_behavior",  "region_distributions.<r>.<dist>.<param>", "Deep override of region distributions"],
            ["charging_infra", "charger_count, *_rate_kw, directionality_frac",
                                                                        "Gate D5 reachability"],
            ["sim_window",     "mode, weekdays_only, start, custom_end", "Day loop in sessions.py"],
            ["noise",          "arrival_time_jitter_min, soc_arrival_jitter_pct, profile",
                                                                        "Post-render jitter; tmyx_stochastic = ±5min, ±3%"],
        ],
        widths_in=[1.1, 2.4, 3.5],
        font_size=8.5,
    )

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    build()
